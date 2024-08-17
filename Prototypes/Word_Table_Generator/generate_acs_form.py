import os
from sqlalchemy import ForeignKey
from sqlalchemy.orm import relationship
from sqlalchemy.sql import func

import argparse
import itertools
import re
import os
import io
import zipfile
from docx import Document
from docx.shared import Inches
from PIL import Image
import pandas as pd
import pydot
from docxcompose.composer import Composer

from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Font, Border, Side
from docx.shared import RGBColor
import docx
from docx.oxml.ns import qn



def create_criterion(docx_file, excel_file, sheet_name, search_string):
    # Load the Excel workbook and sheet
    wb = load_workbook(excel_file, data_only=True)
    sheet = wb[sheet_name]

    # Create a new Document
    #doc = Document()
    #doc.add_heading('Excel Table', level=1)

    # Create a table in the Word document
    num_cols = sheet.max_column
    num_rows = sheet.max_row
    #table = doc.add_table(rows=num_rows, cols=num_cols)

    # Find and replace the target text with the HTML table
    for paragraph in docx_file.paragraphs:
        if search_string in paragraph.text:
            # Get rid of th search string
            paragraph.text = paragraph.text.replace(search_string, "")
            # Create a new table from the DataFrame
            table = docx_file.add_table(rows=num_rows, cols=num_cols)
            #table.style = style
            #for i in range(df.shape[0]):
            #    for j in range(df.shape[1]):
            #        if pd.isna(df.iloc[i, j]):
            #            continue
            #        table.cell(i, j).text = str(df.iloc[i, j])

            paragraph.add_run().element.addnext(table._tbl)

    # Define border style
    border_style = {
        'top': 'single',
        'left': 'single',
        'bottom': 'single',
        'right': 'single',
        'insideH': 'single',
        'insideV': 'single'
    }

    # Apply header row formatting
    for col_index, cell in enumerate(sheet[1], start=0):
        word_cell = table.cell(0, col_index)
        word_cell.text = str(cell.value)
        fill_color = get_fill_color(cell)
        apply_fill_color(word_cell, fill_color)
        font_properties = get_font_properties(cell)
        apply_font_properties(word_cell.add_paragraph(), font_properties)
        apply_border(word_cell, border_style)

    # Apply data rows formatting
    for row_index, row in enumerate(sheet.iter_rows(min_row=2, max_row=num_rows, max_col=num_cols), start=1):
        for col_index, cell in enumerate(row):
            word_cell = table.cell(row_index, col_index)
            word_cell.text = str(cell.value)
            fill_color = get_fill_color(cell)
            apply_fill_color(word_cell, fill_color)
            font_properties = get_font_properties(cell)
            apply_font_properties(word_cell.add_paragraph(), font_properties)
            apply_border(word_cell, border_style)


def main():
    # Create the parser
    parser = argparse.ArgumentParser(description="Program to create accreditation word file from input Excel.")

    # Add arguments
    parser.add_argument('-t', '--template', type=str, help='The template file', default='template.docx')
    parser.add_argument('-i', '--input', type=str, help='The input file', default='input.xlsx')
    parser.add_argument('-o', '--output', type=str, help='The output file', default='output.docx')
    
    # Parse the arguments
    args = parser.parse_args()

    # Use the arguments
    print(f"Template file: {args.template}")
    print(f"Input file: {args.input}")
    print(f"Output file: {args.output}")

    template_file = Document(args.template)
    excel_file = args.input

    criterion = [ 'CriterionA', 'CriterionB', 'CriterionC', 'CriterionD', 'CriterionE' ]

    for criteria in criterion:
        #replace_string_with_table(template_file, excel_file, 'MIT.'+criteria, '{{'+criteria+'}}' )
        excel_to_word_table(template_file, excel_file, 'MIT.'+criteria, '{{'+criteria+'}}' )

    template_file.save(args.output)

if __name__ == "__main__":
    main()
